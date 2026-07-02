"""
新闻收集器：使用Playwright访问主要新闻网站，整理最近一周的十件大事
"""
import asyncio
from datetime import datetime
from pathlib import Path
from playwright.async_api import async_playwright
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 创建输出目录
OUTPUT_DIR = Path("news_output")
OUTPUT_DIR.mkdir(exist_ok=True)

# 新闻网站列表
NEWS_SITES = [
    {
        "name": "新华网",
        "url": "https://www.news.cn/",
        "selectors": {
            "news_list": ".news-list li, .tit-list li, h3 a, .news_item a",
            "title": "a"
        }
    },
    {
        "name": "人民网",
        "url": "https://www.people.com.cn/",
        "selectors": {
            "news_list": ".ej_list_box li, .news-list li, h3 a",
            "title": "a"
        }
    },
    {
        "name": "央视网",
        "url": "https://news.cctv.com/",
        "selectors": {
            "news_list": ".news_list li, .content_list li, h3 a",
            "title": "a"
        }
    }
]


async def scrape_news(page, site_config):
    """从单个网站抓取新闻"""
    news_items = []
    try:
        print(f"正在访问 {site_config['name']}...")
        await page.goto(site_config["url"], wait_until="domcontentloaded", timeout=30000)
        await page.wait_for_timeout(3000)

        # 截图整个页面
        screenshot_path = OUTPUT_DIR / f"{site_config['name']}_homepage.png"
        await page.screenshot(path=str(screenshot_path), full_page=False)
        print(f"  已截图: {screenshot_path}")

        # 尝试提取新闻标题
        selectors = site_config["selectors"]["news_list"].split(", ")

        for selector in selectors:
            try:
                elements = await page.query_selector_all(selector)
                for elem in elements[:15]:  # 每个选择器最多取15条
                    title = await elem.inner_text()
                    title = title.strip()
                    if title and len(title) > 5 and len(title) < 100:
                        link = await elem.get_attribute("href")
                        if not link:
                            link_elem = await elem.query_selector("a")
                            if link_elem:
                                link = await link_elem.get_attribute("href")
                        news_items.append({
                            "title": title,
                            "link": link or "",
                            "source": site_config["name"]
                        })
            except Exception as e:
                continue

        # 去重
        seen = set()
        unique_news = []
        for item in news_items:
            if item["title"] not in seen:
                seen.add(item["title"])
                unique_news.append(item)
        news_items = unique_news[:10]

        print(f"  从 {site_config['name']} 获取到 {len(news_items)} 条新闻")
    except Exception as e:
        print(f"  访问 {site_config['name']} 时出错: {e}")

    return news_items


async def capture_news_detail(page, url, source_name, index):
    """截取新闻详情页"""
    try:
        await page.goto(url, wait_until="domcontentloaded", timeout=20000)
        await page.wait_for_timeout(2000)
        screenshot_path = OUTPUT_DIR / f"news_{index+1}_{source_name}.png"
        await page.screenshot(path=str(screenshot_path), full_page=False)
        return str(screenshot_path)
    except:
        return None


def create_docx(news_items, screenshots):
    """生成Word文档"""
    doc = Document()

    # 标题
    title = doc.add_heading("最近一周国内外大事汇总", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 副标题
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f"整理日期：{datetime.now().strftime('%Y年%m月%d日')}")
    run.font.size = Pt(12)

    doc.add_paragraph()

    # 新闻列表
    for i, item in enumerate(news_items):
        # 新闻标题
        heading = doc.add_heading(f"{i+1}. {item['title']}", level=2)

        # 来源信息
        source_para = doc.add_paragraph()
        run = source_para.add_run(f"来源：{item['source']}")
        run.font.size = Pt(10)
        run.font.color.rgb = None  # 默认颜色

        if item.get('link'):
            link_para = doc.add_paragraph()
            run = link_para.add_run(f"链接：{item['link']}")
            run.font.size = Pt(9)

        # 插入截图
        if i < len(screenshots) and screenshots[i]:
            try:
                doc.add_picture(screenshots[i], width=Inches(5.5))
            except:
                pass

        doc.add_paragraph()  # 空行分隔

    # 保存文档
    doc_path = OUTPUT_DIR / "最近一周大事汇总.docx"
    doc.save(str(doc_path))
    return doc_path


async def main():
    """主函数"""
    print("=" * 50)
    print("新闻收集器 - 整理最近一周的十件大事")
    print("=" * 50)

    all_news = []
    screenshots = []

    async with async_playwright() as p:
        browser = await p.chromium.launch(headless=True)
        context = await browser.new_context(
            viewport={"width": 1280, "height": 800},
            locale="zh-CN"
        )
        page = await context.new_page()

        # 从各个网站收集新闻
        for site in NEWS_SITES:
            news = await scrape_news(page, site)
            all_news.extend(news)

        # 去重并取前10条
        seen = set()
        unique_news = []
        for item in all_news:
            if item['title'] not in seen:
                seen.add(item['title'])
                unique_news.append(item)

        final_news = unique_news[:10]
        print(f"\n共收集到 {len(final_news)} 条重要新闻")

        # 截取新闻详情页
        print("\n正在截取新闻详情页...")
        for i, item in enumerate(final_news):
            if item.get('link'):
                screenshot = await capture_news_detail(
                    page, item['link'], item['source'], i
                )
                screenshots.append(screenshot)
            else:
                screenshots.append(None)

        await browser.close()

    # 生成Word文档
    print("\n正在生成Word文档...")
    doc_path = create_docx(final_news, screenshots)
    print(f"文档已生成: {doc_path}")

    # 打印新闻列表
    print("\n" + "=" * 50)
    print("收集到的十件大事：")
    print("=" * 50)
    for i, item in enumerate(final_news):
        print(f"{i+1}. [{item['source']}] {item['title']}")

    return final_news


if __name__ == "__main__":
    asyncio.run(main())
