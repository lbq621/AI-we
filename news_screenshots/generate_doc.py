"""
生成最近一周大事汇总 Word 文档
"""
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

# 创建文档
doc = Document()

# 标题
title = doc.add_heading('最近一周国内外大事汇总', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 副标题
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run(f'整理日期：{datetime.now().strftime("%Y年%m月%d日")}')
run.font.size = Pt(12)

doc.add_paragraph()

# 十件大事列表
events = [
    {
        "title": "梅西世界杯首秀上演帽子戏法，成世界杯历史射手王",
        "source": "知乎热榜 / 微博热搜",
        "content": "阿根廷3-0完胜阿尔及利亚，梅西上演帽子戏法，成为世界杯历史射手王。梅西与中国球迷的双向奔赴成为热门话题。同时，梅西首秀疑似亮鞋钉踩踏对手小腿的动作引发争议。",
        "screenshot": "zhihu.png"
    },
    {
        "title": "山姆被约谈后更换董事长，阿里系高管接掌",
        "source": "36氪 / 知乎热榜",
        "content": "山姆会员商店中国被监管约谈后，CMO张青提交辞呈。新任董事长由刘鹏接任，引发市场对阿里系高管接掌山姆深层逻辑的讨论。",
        "screenshot": "36kr.png"
    },
    {
        "title": "微信支付推出AI专属卡，支持Agent闭环消费",
        "source": "AIbase / 微博热搜",
        "content": "微信支付正式发布AI专属卡，支持Agent闭环消费，主账户完全隔离。这是国内首个支持AI Agent自主消费的支付产品。",
        "screenshot": "weibo.png"
    },
    {
        "title": "豆包日活超2亿但收入不足百万，字节AI资源或向企业端倾斜",
        "source": "AIbase / 微博热搜",
        "content": "字节跳动旗下AI产品豆包日活用户超过2亿，但每天收入不足百万。消息称字节AI资源可能向企业端倾斜。",
        "screenshot": "weibo.png"
    },
    {
        "title": "必胜客被卖，百胜中国12亿美元收购中国大陆品牌所有权",
        "source": "知乎热榜",
        "content": "百胜中国以12亿美元收购必胜客中国大陆品牌所有权，引发市场对其中国市场前景的讨论。",
        "screenshot": "zhihu.png"
    },
    {
        "title": "智谱AI正式开源GLM-5.2模型",
        "source": "AIbase",
        "content": "智谱AI正式开源GLM-5.2模型，主打1M无损上下文与长程代码任务，为国内AI开源生态再添重磅产品。",
        "screenshot": "36kr.png"
    },
    {
        "title": "小米发布MiMo Claw正式版",
        "source": "AIbase",
        "content": "小米发布MiMo Claw正式版，支持千次连续工具调用，免费时长增至4小时，进一步降低AI使用门槛。",
        "screenshot": "36kr.png"
    },
    {
        "title": "鸿蒙智行全面引入电池二供，国轩高科将上车问界",
        "source": "36氪",
        "content": "鸿蒙智行开始全面引入多家动力电池供应商，国轩高科将上车问界，标志着新能源汽车供应链多元化趋势。",
        "screenshot": "36kr.png"
    },
    {
        "title": "仙工智能港股IPO，市值超112亿港元",
        "source": "36氪",
        "content": "浙大系创业者创办的智能机器人公司仙工智能开始招股，招股价101.60港元，市值112.27亿港元。",
        "screenshot": "36kr.png"
    },
    {
        "title": "SpaceX超越亚马逊跻身全球前五",
        "source": "36氪",
        "content": "SpaceX股价连涨3天，超越亚马逊跻身全球市值前五。美国零售经纪商处的SpaceX投资者每人至少分得一股IPO股票。",
        "screenshot": "36kr.png"
    }
]

# 写入十件大事
doc.add_heading('一、最近一周十件大事', level=1)

for i, event in enumerate(events, 1):
    # 事件标题
    doc.add_heading(f'{i}. {event["title"]}', level=2)

    # 来源
    source_para = doc.add_paragraph()
    run = source_para.add_run(f'来源：{event["source"]}')
    run.font.size = Pt(10)

    # 内容
    content_para = doc.add_paragraph()
    run = content_para.add_run(event["content"])
    run.font.size = Pt(11)

    # 插入截图
    try:
        doc.add_picture(f'D:/AI/we/news_screenshots/{event["screenshot"]}', width=Inches(5.5))
    except Exception as e:
        print(f"无法插入截图 {event['screenshot']}: {e}")

    doc.add_paragraph()  # 空行分隔

# 添加总结
doc.add_heading('二、总结', level=1)
summary = doc.add_paragraph()
summary.add_run('以上是最近一周国内外发生的十件大事，涵盖体育、科技、商业等多个领域。从梅西世界杯创纪录到国内AI产品密集发布，从山姆高层变动到新能源汽车供应链调整，这些事件反映了当前社会的热点话题和发展趋势。')

# 保存文档
output_path = 'D:/AI/we/news_screenshots/最近一周大事汇总.docx'
doc.save(output_path)
print(f'文档已生成: {output_path}')
